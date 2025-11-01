from flask import Blueprint, render_template, request, flash, send_file
from textblob import TextBlob
import os
import docx
from PyPDF2 import PdfReader
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from collections import Counter
from datetime import datetime


report_bp = Blueprint('report', __name__, template_folder='templates', static_folder='static')

@report_bp.route('/report', methods=['GET', 'POST'])
def analyze_report():
    if request.method == 'POST':
        file = request.files['document']
        if not file:
            flash("No file uploaded!", "danger")
            return render_template('report.html')

        text = extract_text(file)
        if not text:
            flash("Unsupported file format or empty document.", "danger")
            return render_template('report.html')
        
        meta = extract_metadata(file)

        blob = TextBlob(text)
        analysis = {
            "summary": "Document analyzed successfully.",
            "sentiment": {
                "polarity": round(blob.sentiment.polarity, 3),
                "subjectivity": round(blob.sentiment.subjectivity, 3)
            },
            "top_words": Counter(blob.word_counts).most_common(5)

        }

        pdf_path = os.path.join('static', 'reports', 'report.pdf')
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        generate_pdf_report(analysis, pdf_path)

        flash("Report generated successfully!", "success")
        return render_template('report.html', analysis=analysis, report_url=pdf_path, metadata=meta)

    return render_template('report.html')


def extract_text(file):
    if file.filename.endswith('.txt'):
        return file.read().decode('utf-8')
    elif file.filename.endswith('.pdf'):
        reader = PdfReader(file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file.filename.endswith('.docx'):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return None


def generate_pdf_report(analysis, output_path):
    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Document Report", styles['Title']),
        Paragraph(f"Sentiment Polarity: {analysis['sentiment']['polarity']}", styles['Normal']),
        Paragraph(f"Sentiment Subjectivity: {analysis['sentiment']['subjectivity']}", styles['Normal']),
        Paragraph("Top Words:", styles['Heading2']),
    ]
    for word, count in analysis['top_words']:
        story.append(Paragraph(f"{word}: {count}", styles['Normal']))
    doc.build(story)

from datetime import datetime
import os
import docx
from PyPDF2 import PdfReader

def extract_metadata(file):
    metadata = {
        "File Name": file.filename,
        "Upload Time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

    # Get file creation/modification time (from your local filesystem, if saved)
    try:
        # If file has been saved temporarily
        temp_path = os.path.join("uploads", file.filename)
        os.makedirs("uploads", exist_ok=True)
        file.save(temp_path)

        creation_time = os.path.getctime(temp_path)
        modified_time = os.path.getmtime(temp_path)
        metadata["File Created (Local)"] = datetime.fromtimestamp(creation_time).strftime("%Y-%m-%d %H:%M:%S")
        metadata["File Modified (Local)"] = datetime.fromtimestamp(modified_time).strftime("%Y-%m-%d %H:%M:%S")
    except Exception as e:
        metadata["File Created (Local)"] = "Unavailable"

    # Reset pointer for format-specific metadata reading
    file.stream.seek(0)

    # PDF Metadata
    if file.filename.endswith('.pdf'):
        reader = PdfReader(file)
        info = reader.metadata
        if info:
            metadata.update({
                "Author": info.author or "Unknown",
                "Creator": info.creator or "Unknown",
                "Producer": info.producer or "Unknown",
                "Pages": len(reader.pages)
            })
            # PDF creation date (usually stored as D:YYYYMMDDHHMMSS)
            if info.get('/CreationDate'):
                raw_date = info.get('/CreationDate').replace("D:", "").split('+')[0]
                try:
                    dt = datetime.strptime(raw_date, "%Y%m%d%H%M%S")
                    metadata["Document Created (From PDF)"] = dt.strftime("%Y-%m-%d %H:%M:%S")
                except Exception:
                    metadata["Document Created (From PDF)"] = raw_date

    # DOCX Metadata
    elif file.filename.endswith('.docx'):
        doc = docx.Document(file)
        props = doc.core_properties
        metadata.update({
            "Author": props.author or "Unknown",
            "Title": props.title or "Untitled",
            "Last Modified": props.modified.strftime("%Y-%m-%d %H:%M:%S") if props.modified else "Unknown",
            "Created (From DOCX)": props.created.strftime("%Y-%m-%d %H:%M:%S") if props.created else "Unknown",
        })

    # TXT Metadata
    elif file.filename.endswith('.txt'):
        metadata.update({
            "File Type": "Plain Text (.txt)"
        })

    return metadata
